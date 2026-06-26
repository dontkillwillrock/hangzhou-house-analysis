"""将项目计划书.md转换为Word文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def add_table_from_rows(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D9E2F3')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行


def build_document():
    doc = Document()

    # ── 全局样式 ──
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25

    # 标题样式
    for level, size, color in [(1, 22, '1F4E79'), (2, 16, '2E75B6'), (3, 13, '2E75B6')]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '黑体'
        hs.font.size = Pt(size)
        hs.font.color.rgb = RGBColor(*bytes.fromhex(color))
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # ═══════════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('杭州学区房数据分析\n项目计划书')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('杭州学区房数据分析与房价预测')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for _ in range(6):
        doc.add_paragraph()

    info_lines = [
        ('指导教师：', '________'),
        ('团队人数：', '4人'),
        ('项目周期：', '4.5小时'),
        ('编制日期：', '2026年6月'),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}{value}')
        run.font.size = Pt(13)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 目录页（手动占位）
    # ═══════════════════════════════════════════
    toc_title = doc.add_heading('目  录', level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_items = [
        '一、项目基本信息 ……………………………………… 3',
        '二、团队分工 …………………………………………… 3',
        '三、项目目标 …………………………………………… 4',
        '四、技术方案 …………………………………………… 5',
        '五、项目计划 …………………………………………… 6',
        '六、数据说明 …………………………………………… 8',
        '七、预期结果 …………………………………………… 9',
        '八、风险管理 …………………………………………… 9',
        '九、提交材料 …………………………………………… 10',
        '十、评分标准 …………………………………………… 10',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 一、项目基本信息
    # ═══════════════════════════════════════════
    doc.add_heading('一、项目基本信息', level=1)
    add_table_from_rows(doc,
        ['项目', '内容'],
        [
            ['项目名称', '杭州学区房数据分析与房价预测'],
            ['项目类型', '数据分析与机器学习'],
            ['团队人数', '4人'],
            ['项目周期', '4.5小时（3个阶段，每阶段90分钟）'],
            ['指导教师', '________'],
            ['小组成员', '见下表'],
        ],
        col_widths=[4, 10],
    )

    # ═══════════════════════════════════════════
    # 二、团队分工
    # ═══════════════════════════════════════════
    doc.add_heading('二、团队分工', level=1)
    add_table_from_rows(doc,
        ['角色', '姓名', '学号', '主要职责'],
        [
            ['组长', '苏毅', '________', '项目统筹、进度管理、任务分配、报告审核'],
            ['组员1', '李文锋', '________', '阶段一：数据清洗与特征提取'],
            ['组员2', '梁俊杰', '________', '阶段二：模型训练与交叉验证'],
            ['组员3', '林文彬', '________', '阶段三：模型优化与报告撰写'],
        ],
        col_widths=[2.5, 2.5, 3, 6],
    )

    doc.add_heading('组长职责', level=2)
    duties = [
        '制定项目计划，分配任务',
        '监督各阶段进度，协调组员工作',
        '审核代码质量和实验结果',
        '组织团队讨论，解决技术问题',
        '最终报告的汇总和提交',
    ]
    for i, d in enumerate(duties, 1):
        doc.add_paragraph(f'{i}. {d}')

    # ═══════════════════════════════════════════
    # 三、项目目标
    # ═══════════════════════════════════════════
    doc.add_heading('三、项目目标', level=1)

    doc.add_heading('3.1 业务目标', level=2)
    doc.add_paragraph(
        '预测杭州学区房的单价，为购房者提供价格参考，帮助评估学区房价格是否合理，辅助投资决策。'
    )

    doc.add_heading('3.2 技术目标', level=2)
    tech_goals = [
        '完成数据探索性分析（EDA）',
        '构建房价预测回归模型',
        '分析影响房价的关键因素',
        '实现模型融合提升预测精度',
    ]
    for g in tech_goals:
        doc.add_paragraph(g, style='List Bullet')

    doc.add_heading('3.3 预期成果', level=2)
    results = ['完整的分析报告', '可复用的代码模块', '训练好的预测模型', '业务洞察与建议']
    for r in results:
        doc.add_paragraph(r, style='List Bullet')

    # ═══════════════════════════════════════════
    # 四、技术方案
    # ═══════════════════════════════════════════
    doc.add_heading('四、技术方案', level=1)

    doc.add_heading('4.1 技术栈', level=2)
    add_table_from_rows(doc,
        ['类别', '技术/工具'],
        [
            ['编程语言', 'Python 3.9+'],
            ['数据处理', 'pandas, numpy'],
            ['机器学习', 'scikit-learn'],
            ['可视化', 'matplotlib, seaborn'],
            ['开发环境', 'Jupyter Notebook / VS Code'],
        ],
        col_widths=[4, 10],
    )

    doc.add_heading('4.2 算法选型', level=2)
    add_table_from_rows(doc,
        ['模型', '类型', '用途'],
        [
            ['线性回归', '线性模型', '基线模型'],
            ['随机森林', 'Bagging集成', '主力模型'],
            ['梯度提升(GBDT)', 'Boosting集成', '主力模型'],
            ['Stacking融合', '模型融合', '最终模型'],
        ],
        col_widths=[4, 4, 4],
    )

    doc.add_heading('4.3 评估指标', level=2)
    add_table_from_rows(doc,
        ['指标', '说明', '目标值'],
        [
            ['R²', '决定系数', '> 0.80'],
            ['RMSE', '均方根误差', '越小越好'],
            ['MAE', '平均绝对误差', '越小越好'],
        ],
        col_widths=[3, 5, 4],
    )

    # ═══════════════════════════════════════════
    # 五、项目计划
    # ═══════════════════════════════════════════
    doc.add_heading('五、项目计划', level=1)

    doc.add_heading('5.1 时间安排', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '项目总时长4.5小时，分为三个阶段，每阶段90分钟：'
    )
    run.font.size = Pt(11)

    # 用表格模拟时间线
    timeline = doc.add_table(rows=2, cols=3)
    timeline.style = 'Table Grid'
    timeline.alignment = WD_TABLE_ALIGNMENT.CENTER
    phases = [
        ('阶段一（90分钟）', '数据探索\n数据清洗\n基础特征', '李文锋+苏毅'),
        ('阶段二（90分钟）', '特征工程\n模型训练\n交叉验证', '梁俊杰'),
        ('阶段三（90分钟）', '模型优化\n报告撰写\n成果汇总', '林文彬+苏毅'),
    ]
    for i, (title, content, owner) in enumerate(phases):
        cell = timeline.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')

        cell2 = timeline.rows[1].cells[i]
        cell2.text = f'{content}\n负责人：{owner}'
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell2.paragraphs[0].runs:
            run.font.size = Pt(9)
    doc.add_paragraph()

    # 阶段一
    doc.add_heading('阶段一：数据获取、探索分析与特征工程基础（90分钟）', level=3)
    add_table_from_rows(doc,
        ['时间', '任务', '负责人', '产出'],
        [
            ['0-10min', '理解业务问题与字段角色', '全员', '业务理解文档'],
            ['10-25min', '读取数据并检查结构', '李文锋', '数据结构报告'],
            ['25-40min', '数据质量评估', '李文锋', '质量评估表'],
            ['40-55min', '基础统计描述', '李文锋+苏毅', '统计描述表'],
            ['55-75min', '数据清洗与特征提取', '李文锋', '清洗后数据'],
            ['75-90min', '阶段总结与检查', '苏毅', '阶段报告'],
        ],
        col_widths=[2.5, 5, 3, 3],
    )

    # 阶段二
    doc.add_heading('阶段二：深度特征工程与多模型训练（90分钟）', level=3)
    add_table_from_rows(doc,
        ['时间', '任务', '负责人', '产出'],
        [
            ['0-10min', '回顾与规划', '梁俊杰', '特征工程计划'],
            ['10-30min', '深度特征工程', '梁俊杰', '衍生特征'],
            ['30-45min', '系统性可视化分析', '梁俊杰', '可视化图表'],
            ['45-55min', '类别编码与数据准备', '梁俊杰', '建模数据'],
            ['55-75min', '多模型训练', '梁俊杰', '训练好的模型'],
            ['75-90min', '交叉验证与对比', '梁俊杰+苏毅', '模型对比表'],
        ],
        col_widths=[2.5, 5, 3, 3],
    )

    # 阶段三
    doc.add_heading('阶段三：迭代优化与模型评估（90分钟）', level=3)
    add_table_from_rows(doc,
        ['时间', '任务', '负责人', '产出'],
        [
            ['0-25min', '超参调优', '林文彬', '调优后模型'],
            ['25-40min', '模型融合与集成', '林文彬', '融合模型'],
            ['40-55min', '模型评估与可视化', '林文彬', '评估图表'],
            ['55-70min', '特征重要性与业务洞察', '林文彬+苏毅', '业务建议'],
            ['70-85min', '实验报告撰写', '全员', '实验报告'],
            ['85-90min', '总结与提交', '苏毅', '最终提交包'],
        ],
        col_widths=[2.5, 5, 3, 3],
    )

    # ═══════════════════════════════════════════
    # 六、数据说明
    # ═══════════════════════════════════════════
    doc.add_heading('六、数据说明', level=1)

    doc.add_heading('6.1 数据来源', level=2)
    doc.add_paragraph('杭州二手房交易数据，包含学区房相关信息。')

    doc.add_heading('6.2 数据规模', level=2)
    for item in ['样本数量：9200条', '特征数量：26个', '目标变量：单价（元/平米）']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.3 主要字段', level=2)
    add_table_from_rows(doc,
        ['字段', '类型', '说明'],
        [
            ['建筑面积', '数值', '房屋面积（平方米）'],
            ['户型', '文本', '如3室2厅'],
            ['楼层', '类别', '低层/中层/高层'],
            ['朝向', '类别', '南/南北/东南等'],
            ['建筑年代', '数值', '建造年份'],
            ['单价（元/平米）', '数值', '目标变量'],
            ['总价（万元）', '数值', '房屋总价'],
            ['所在学校', '文本', '对应学校'],
            ['距离学校', '文本', '与学校距离'],
            ['所在区县', '类别', '区域位置'],
        ],
        col_widths=[3.5, 2.5, 7],
    )

    # ═══════════════════════════════════════════
    # 七、预期结果
    # ═══════════════════════════════════════════
    doc.add_heading('七、预期结果', level=1)

    doc.add_heading('7.1 模型性能', level=2)
    add_table_from_rows(doc,
        ['模型', '预期R²', '实际R²'],
        [
            ['线性回归（基线）', '0.50-0.60', '0.5474'],
            ['随机森林', '0.75-0.85', '0.8411'],
            ['梯度提升', '0.75-0.85', '0.8264'],
            ['Stacking融合', '0.80-0.90', '0.8409'],
        ],
        col_widths=[5, 3, 3],
    )

    doc.add_heading('7.2 业务洞察', level=2)
    insights = [
        '建筑面积是影响房价的最重要因素',
        '学校等级评分对房价有显著影响',
        '距离学校越近，房价越高',
        '房龄影响房价，次新房价格最高',
        '楼层对房价有影响，高层和中层更受欢迎',
    ]
    for i, ins in enumerate(insights, 1):
        doc.add_paragraph(f'{i}. {ins}')

    # ═══════════════════════════════════════════
    # 八、风险管理
    # ═══════════════════════════════════════════
    doc.add_heading('八、风险管理', level=1)
    add_table_from_rows(doc,
        ['风险', '影响', '应对措施'],
        [
            ['数据质量问题', '模型效果差', '严格数据清洗，多重验证'],
            ['模型过拟合', '泛化能力差', '交叉验证，正则化'],
            ['时间不足', '任务未完成', '优先核心任务，合理分工'],
            ['技术难题', '进度延误', '及时讨论，寻求帮助'],
        ],
        col_widths=[4, 3.5, 5.5],
    )

    # ═══════════════════════════════════════════
    # 九、提交材料
    # ═══════════════════════════════════════════
    doc.add_heading('九、提交材料', level=1)

    doc.add_heading('9.1 代码文件', level=2)
    code_items = [
        'experiment_06/notebooks/ — 阶段一～三完整代码',
        'experiment_06/src/ — 可复用模块',
        'experiment_06/models/ — 训练好的模型',
        'experiment_06/figures/ — 可视化图表',
        'experiment_06/reports/ — 实验报告',
    ]
    for item in code_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.2 文档材料', level=2)
    docs_list = ['项目计划书', '实验报告', '代码说明文档', '数据分析报告']
    for item in docs_list:
        doc.add_paragraph(f'☐ {item}')

    doc.add_heading('9.3 演示材料', level=2)
    demo_list = ['PPT演示文稿', '项目演示视频（可选）']
    for item in demo_list:
        doc.add_paragraph(f'☐ {item}')

    # ═══════════════════════════════════════════
    # 十、评分标准
    # ═══════════════════════════════════════════
    doc.add_heading('十、评分标准', level=1)
    add_table_from_rows(doc,
        ['评分项', '分值', '说明'],
        [
            ['数据清洗与特征工程', '25%', '缺失值处理、特征构造'],
            ['模型训练与评估', '30%', '模型选择、交叉验证'],
            ['模型优化与融合', '20%', '超参调优、模型融合'],
            ['可视化与分析', '15%', '图表质量、业务洞察'],
            ['报告与演示', '10%', '报告规范、表达清晰'],
        ],
        col_widths=[4.5, 2.5, 6],
    )

    # ═══════════════════════════════════════════
    # 签名页
    # ═══════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('项目计划书编制：组长 苏毅')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('编制日期：2026年6月')
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_heading('组员签名', level=2)
    members = ['李文锋', '梁俊杰', '林文彬']
    for m in members:
        p = doc.add_paragraph()
        run = p.add_run(f'{m}  ________________')
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(16)

    # ── 保存 ──
    out = r'f:\zuoye\数据分析\实验6\experiment_06\项目计划书.docx'
    doc.save(out)
    print(f'✓ 已生成: {out}')


if __name__ == '__main__':
    build_document()
